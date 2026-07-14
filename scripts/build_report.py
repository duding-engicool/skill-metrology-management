# -*- coding: utf-8 -*-
"""
计量管理报告渲染器

功能：依据"上次校准日期 + 校准周期(月)"自动计算下次校准日期，并做三级超期预警，
      默认输出 Markdown 台账报告；可按需导出 Word(.docx) / Excel(.xlsx)。

用法：
  python build_report.py                              # 内置样本 → 仅 MD（默认）
  python build_report.py --data-file instruments.json # 自定义台账
  python build_report.py --format all                 # MD + DOCX + XLSX
  python build_report.py --format xlsx --out-dir ./out# 仅 Excel
  python build_report.py --as-of 2026-07-13           # 指定基准日期

说明：
  - 默认仅依赖标准库输出 MD；DOCX/XLSX 在对应 --format 时才惰性导入
    openpyxl / python-docx（需 pip install openpyxl python-docx）。
  - 输出默认写到【用户当前工作目录】，可用 --out-dir 指定；技能目录本身不含产物。
"""
import argparse
import json
import os
from datetime import date

PRIMARY = "#C8102E"
STATUS_COLOR = {
    "超期(红色)": "#C8102E",
    "预警(橙色)": "#E8741C",
    "预警(黄色)": "#D9A400",
    "正常": "#2E7D32",
}

# 内置小样本：5 台器具，覆盖 正常/超期/超期(强检)/橙色/黄色 四态
SAMPLE = {
    "instruments": [
        {"name": "数显卡尺 0-150mm", "code": "CL-001", "category": "长度",
         "last_cal": "2026-01-15", "cycle_months": 12, "mandatory": False, "cert_no": "待企业补充"},
        {"name": "千分尺 0-25mm", "code": "CL-002", "category": "长度",
         "last_cal": "2025-06-20", "cycle_months": 12, "mandatory": False, "cert_no": "待企业补充"},
        {"name": "一般压力表 0-1MPa(强检)", "code": "CL-003", "category": "压力",
         "last_cal": "2025-12-01", "cycle_months": 6, "mandatory": True, "cert_no": "待企业补充"},
        {"name": "扭矩扳手 5-50N·m", "code": "CL-004", "category": "力学",
         "last_cal": "2026-01-15", "cycle_months": 6, "mandatory": False, "cert_no": "待企业补充"},
        {"name": "工作用热电偶 K型", "code": "CL-005", "category": "温度",
         "last_cal": "2025-07-30", "cycle_months": 12, "mandatory": False, "cert_no": "待企业补充"},
    ]
}


def parse_date(s):
    return date.fromisoformat(s)


def add_months(d, months):
    """日期加月份，日取月末兜底。"""
    m = d.month - 1 + months
    y = d.year + m // 12
    m = m % 12 + 1
    day = min(d.day, [31, 29 if (y % 4 == 0 and y % 100 != 0) or y % 400 == 0 else 28,
                      31, 30, 31, 30, 31, 31, 30, 31, 30, 31][m - 1])
    return date(y, m, day)


def evaluate(inst, as_of):
    """计算下次校准日期与预警状态。"""
    last = parse_date(inst["last_cal"])
    cycle = int(inst.get("cycle_months", 12))
    nxt = add_months(last, cycle)
    delta = (nxt - as_of).days
    if nxt < as_of:
        status, advise = "超期(红色)", "立即停用，安排送检"
        if inst.get("mandatory"):
            advise = "强检设备超期，判定不合格，安排检定"
    elif delta <= 7:
        status, advise = "预警(橙色)", "本周内必须送检"
    elif delta <= 30:
        status, advise = "预警(黄色)", "尽快安排送检"
    else:
        status, advise = "正常", "保持监测"
    return nxt.isoformat(), status, advise


def build_rows(data, as_of):
    rows = []
    counts = {"超期(红色)": 0, "预警(橙色)": 0, "预警(黄色)": 0, "正常": 0}
    for it in data["instruments"]:
        nxt, status, advise = evaluate(it, as_of)
        counts[status] = counts.get(status, 0) + 1
        mand = "（强检）" if it.get("mandatory") and "强检" not in it["name"] else ""
        rows.append({
            "name": it["name"] + mand,
            "code": it["code"],
            "category": it["category"],
            "last_cal": it["last_cal"],
            "cycle": it["cycle_months"],
            "next": nxt,
            "status": status,
            "advise": advise,
        })
    return rows, counts


def build_md(data, as_of):
    rows, counts = build_rows(data, as_of)
    lines = []
    lines.append("# 计量器具台账与校准状态报告")
    lines.append("")
    lines.append(f"> 基准日期：{as_of}　器具数量：{len(rows)}")
    lines.append("")
    lines.append("| 器具名称 | 编号 | 类别 | 上次校准 | 周期(月) | 下次校准 | 状态 | 送检建议 |")
    lines.append("|----------|------|------|----------|----------|----------|------|----------|")
    for r in rows:
        lines.append(f"| {r['name']} | {r['code']} | {r['category']} | "
                     f"{r['last_cal']} | {r['cycle']} | {r['next']} | {r['status']} | {r['advise']} |")
    lines.append("")
    lines.append("## 预警汇总")
    lines.append("")
    lines.append(f"- 超期(红)：{counts['超期(红色)']} 台　预警(橙)：{counts['预警(橙色)']} 台　"
                 f"预警(黄)：{counts['预警(黄色)']} 台　正常：{counts['正常']} 台")
    lines.append("")
    lines.append("> 合规红线：超期器具必须立即停用；强检设备超期直接判定不合格。证书结论以实际为准。")
    lines.append("")
    lines.append("---")
    lines.append(f"*本报告由 metrology-management 自动生成，默认 MD 台账；如需 Word/Excel 版本可另行导出*")
    return "\n".join(lines)


def build_docx(data, as_of, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows, counts = build_rows(data, as_of)
    doc = Document()
    title = doc.add_heading("计量器具台账与校准状态报告", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)
    p = doc.add_paragraph()
    p.add_run(f"基准日期：{as_of}　器具数量：{len(rows)}").italic = True

    headers = ["器具名称", "编号", "类别", "上次校准", "周期(月)", "下次校准", "状态", "送检建议"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate([r["name"], r["code"], r["category"], r["last_cal"],
                               str(r["cycle"]), r["next"], r["status"], r["advise"]]):
            cells[i].text = str(v)

    doc.add_heading("预警汇总", level=2)
    doc.add_paragraph(
        f"超期(红)：{counts['超期(红色)']} 台　预警(橙)：{counts['预警(橙色)']} 台　"
        f"预警(黄)：{counts['预警(黄色)']} 台　正常：{counts['正常']} 台"
    )
    red = doc.add_paragraph()
    run = red.add_run("合规红线：超期器具必须立即停用；强检设备超期直接判定不合格。证书结论以实际为准。")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)
    doc.save(path)


def build_xlsx(data, as_of, path):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    rows, counts = build_rows(data, as_of)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "计量台账"
    headers = ["器具名称", "编号", "类别", "上次校准", "周期(月)", "下次校准", "状态", "送检建议"]
    header_fill = PatternFill("solid", fgColor="C8102E")
    for c, h in enumerate(headers, 1):
        cell = ws.cell(1, c, h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for i, r in enumerate(rows, 2):
        vals = [r["name"], r["code"], r["category"], r["last_cal"],
                r["cycle"], r["next"], r["status"], r["advise"]]
        for c, v in enumerate(vals, 1):
            ws.cell(i, c, v)
        # 状态列着色
        sc = ws.cell(i, 7)
        hexc = STATUS_COLOR.get(r["status"], "#333333").lstrip("#")
        sc.font = Font(bold=True, color=hexc)
    # 预警汇总
    srow = len(rows) + 3
    ws.cell(srow, 1, "预警汇总").font = Font(bold=True, color="C8102E")
    ws.cell(srow + 1, 1,
            f"超期(红)：{counts['超期(红色)']} 台　预警(橙)：{counts['预警(橙色)']} 台　"
            f"预警(黄)：{counts['预警(黄色)']} 台　正常：{counts['正常']} 台")
    ws.cell(srow + 2, 1, "合规红线：超期器具必须立即停用；强检设备超期直接判定不合格。证书结论以实际为准。")
    ws.column_dimensions["A"].width = 26
    for col in "BCDEFGH":
        ws.column_dimensions[col].width = 14
    wb.save(path)


def main():
    parser = argparse.ArgumentParser(description="计量管理报告（默认 MD，可选 DOCX/XLSX）")
    parser.add_argument("--data-file", default="", help="JSON 台账文件路径")
    parser.add_argument("--out-dir", default="", help="输出目录，默认当前工作目录")
    parser.add_argument("--as-of", default="", help="基准日期 YYYY-MM-DD，默认今天")
    parser.add_argument("--format", default="md",
                        choices=["md", "docx", "xlsx", "all"],
                        help="输出格式：md(默认) / docx / xlsx / all")
    args = parser.parse_args()

    data = SAMPLE
    if args.data_file:
        with open(args.data_file, "r", encoding="utf-8") as f:
            data = json.load(f)
    as_of = parse_date(args.as_of) if args.as_of else date.today()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    out_dir = args.out_dir or os.getcwd()
    os.makedirs(out_dir, exist_ok=True)

    fmt = args.format
    if fmt in ("md", "all"):
        md_path = os.path.join(out_dir, "metrology_report.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(build_md(data, as_of))
        print(f"[OK] MD   -> {os.path.abspath(md_path)}")
    if fmt in ("docx", "all"):
        docx_path = os.path.join(out_dir, "metrology_report.docx")
        build_docx(data, as_of, docx_path)
        print(f"[OK] DOCX -> {os.path.abspath(docx_path)}")
    if fmt in ("xlsx", "all"):
        xlsx_path = os.path.join(out_dir, "metrology_report.xlsx")
        build_xlsx(data, as_of, xlsx_path)
        print(f"[OK] XLSX -> {os.path.abspath(xlsx_path)}")


if __name__ == "__main__":
    main()
